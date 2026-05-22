"""Reconstruct activity — write final DOCX with text substitution.

Phase 3 is DOCX-only by design. We open the original .docx and replace
paragraph/run text with translated_text keyed by segment_id where possible;
fall back to in-order substitution when CU didn't return stable IDs.

Also writes the audit bundle (segments + judge result + metadata) and a
flat translated-text JSON to the `final/` and `audit/` containers.
"""
from __future__ import annotations

import io
import logging
import os
from datetime import datetime, timezone

import azure.durable_functions as df

from app.clients.blob import download_bytes, upload_bytes, upload_json
from app.models import Segment

reconstruct_bp = df.Blueprint()


def _split_blob_path(path: str) -> tuple[str, str]:
    container, _, rest = path.partition("/")
    return (container or "inbound", rest or path)


def _reconstruct_docx(source_bytes: bytes, segments: list[Segment]) -> bytes:
    """Open a docx and substitute paragraph/cell text in document order.

    If the count matches, this preserves formatting (runs, styles, images,
    headers/footers) better than rebuilding from scratch.
    """
    from docx import Document

    doc = Document(io.BytesIO(source_bytes))
    by_id = {s.segment_id: s for s in segments}
    in_order = [s for s in segments if s.translated_text]
    order_idx = 0

    def _replace_para(p):
        nonlocal order_idx
        if not p.text or not p.text.strip():
            return
        translated = None
        # Try ID match first (CU IDs aren't typically embedded in docx, but no harm)
        if p.text in by_id:
            translated = by_id[p.text].translated_text
        if not translated and order_idx < len(in_order):
            translated = in_order[order_idx].translated_text
            order_idx += 1
        if not translated:
            return
        # Replace by clearing runs and writing into the first run to preserve its style
        if p.runs:
            p.runs[0].text = translated
            for r in p.runs[1:]:
                r.text = ""
        else:
            p.text = translated

    for p in doc.paragraphs:
        _replace_para(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_para(p)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


@reconstruct_bp.activity_trigger(input_name="payload")
def activity_reconstruct(payload: dict) -> dict:
    job_id = payload["jobId"]
    target_language = payload["targetLanguage"]
    source_blob = payload["sourceBlob"]
    judge_result = payload["judgeResult"]
    attempts = payload.get("attempts", 1)
    segments = [Segment(**s) for s in payload["segments"]]

    container, blob_path = _split_blob_path(source_blob)
    source_bytes = download_bytes(container, blob_path)
    ext = os.path.splitext(blob_path)[1].lower()

    if ext == ".docx":
        try:
            final_bytes = _reconstruct_docx(source_bytes, segments)
            final_name = f"{job_id}/{target_language}/translated.docx"
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            upload_bytes("final", final_name, final_bytes, content_type=content_type)
            final_blob = f"final/{final_name}"
        except Exception as exc:
            logging.exception("reconstruct jobId=%s docx failed, falling back to txt: %s", job_id, exc)
            final_blob = _write_text_fallback(job_id, target_language, segments)
    else:
        final_blob = _write_text_fallback(job_id, target_language, segments)

    audit_payload = {
        "jobId": job_id,
        "targetLanguage": target_language,
        "sourceBlob": source_blob,
        "finalBlob": final_blob,
        "attempts": attempts,
        "judgeResult": judge_result,
        "segments": [s.model_dump() for s in segments],
        "completedAt": datetime.now(timezone.utc).isoformat(),
    }
    audit_name = f"{job_id}/audit.json"
    upload_json("audit", audit_name, audit_payload)

    logging.info("reconstruct jobId=%s final=%s audit=audit/%s", job_id, final_blob, audit_name)
    return {
        "finalBlob": final_blob,
        "auditBlob": f"audit/{audit_name}",
    }


def _write_text_fallback(job_id: str, target_language: str, segments: list[Segment]) -> str:
    lines = []
    for s in segments:
        if s.kind == "heading":
            lines.append("")
            lines.append(f"# {s.translated_text or s.source_text}")
        elif s.kind == "list_item":
            lines.append(f"- {s.translated_text or s.source_text}")
        else:
            lines.append(s.translated_text or s.source_text)
    payload = "\n".join(lines).encode("utf-8")
    name = f"{job_id}/{target_language}/translated.txt"
    upload_bytes("final", name, payload, content_type="text/plain; charset=utf-8")
    return f"final/{name}"


# Route-to-review handler — also lives here so the orchestrator gets one stop
@reconstruct_bp.activity_trigger(input_name="payload")
def activity_route_to_review(payload: dict) -> dict:
    job_id = payload["jobId"]
    segments = [Segment(**s) for s in payload["segments"]]
    judge_result = payload["judgeResult"]
    attempts = payload["attempts"]

    review_payload = {
        "jobId": job_id,
        "attempts": attempts,
        "judgeResult": judge_result,
        "segments": [s.model_dump() for s in segments],
        "queuedAt": datetime.now(timezone.utc).isoformat(),
    }
    review_name = f"{job_id}/review.json"
    upload_json("reviewed", review_name, review_payload)
    logging.info("route_to_review jobId=%s blob=reviewed/%s", job_id, review_name)
    return {"reviewBlob": f"reviewed/{review_name}"}
