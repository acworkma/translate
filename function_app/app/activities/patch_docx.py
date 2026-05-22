"""Patch DOCX activity — apply revised segment text onto the Document-Translation
output, then finalize.

Document Translation already produced a format-preserving DOCX in `translated/...`.
For segments the LLM did NOT revise, we leave them alone — preserving images,
run-level formatting, table layout, headers, and footers.

For segments that WERE revised by the reviser, we locate the paragraph by
text match against `segment.dt_text` (the Document Translation output that we
snapshotted in `pair_segments`). The first run keeps its formatting and gets
the revised text; remaining runs in the same paragraph are emptied.

Locating by text instead of index sidesteps any ordering drift between
Document Intelligence reading order and python-docx body iteration.
"""
from __future__ import annotations

import io
import logging
import os

import azure.durable_functions as df
from docx import Document
from docx.document import Document as _Document
from docx.text.paragraph import Paragraph

from app.clients.blob import download_bytes, split_path, upload_bytes, upload_json
from app.models import Segment

patch_docx_bp = df.Blueprint()


def _iter_text_units(doc: _Document):
    """Yield every body paragraph plus all paragraphs inside table cells."""
    for para in doc.paragraphs:
        yield para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para


def _set_paragraph_text(para: Paragraph, new_text: str) -> None:
    runs = para.runs
    if not runs:
        para.add_run(new_text)
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""


@patch_docx_bp.activity_trigger(input_name="payload")
def activity_patch_docx(payload: dict) -> dict:
    job_id = payload["jobId"]
    target_language = payload["targetLanguage"]
    translated_blob = payload["translatedBlob"]
    segments = [Segment(**s) for s in payload["segments"]]

    container, path = split_path(translated_blob, default_container="translated")
    docx_bytes = download_bytes(container, path)
    doc = Document(io.BytesIO(docx_bytes))

    revised = [
        s for s in segments
        if s.dt_text and s.translated_text and s.translated_text.strip() != (s.dt_text or "").strip()
    ]

    patched = 0
    unresolved: list[str] = []

    if revised:
        units = list(_iter_text_units(doc))
        consumed: set[int] = set()
        by_text: dict[str, list[int]] = {}
        for i, p in enumerate(units):
            t = (p.text or "").strip()
            if t:
                by_text.setdefault(t, []).append(i)

        for seg in revised:
            key = (seg.dt_text or "").strip()
            indexes = by_text.get(key, [])
            target_idx = next((i for i in indexes if i not in consumed), None)
            if target_idx is None:
                unresolved.append(seg.segment_id)
                continue
            _set_paragraph_text(units[target_idx], seg.translated_text)
            consumed.add(target_idx)
            patched += 1

    if unresolved:
        logging.warning(
            "patch_docx jobId=%s unresolved revised segments: %s",
            job_id, unresolved[:10],
        )

    buf = io.BytesIO()
    doc.save(buf)
    final_bytes = buf.getvalue()

    filename = os.path.basename(path)
    final_path = f"{job_id}/{target_language}/{filename}"
    upload_bytes(
        "final",
        final_path,
        final_bytes,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    audit = {
        "jobId": job_id,
        "targetLanguage": target_language,
        "sourceBlob": payload.get("sourceBlob"),
        "translatedBlob": translated_blob,
        "finalBlob": f"final/{final_path}",
        "attempts": payload.get("attempts"),
        "judgeResult": payload.get("judgeResult"),
        "revisedSegments": len(revised),
        "patchedSegments": patched,
        "unresolvedRevised": unresolved,
        "totalSegments": len(segments),
        "segments": [s.model_dump() for s in segments],
    }
    audit_path = f"{job_id}/audit.json"
    upload_json("audit", audit_path, audit)

    logging.info(
        "patch_docx jobId=%s revised=%d patched=%d unresolved=%d final=%s",
        job_id, len(revised), patched, len(unresolved), final_path,
    )
    return {
        "finalBlob": f"final/{final_path}",
        "auditBlob": f"audit/{audit_path}",
        "patchedSegments": patched,
        "revisedSegments": len(revised),
    }
