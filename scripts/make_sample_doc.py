"""Generate a small synthetic discharge-summary DOCX for smoke testing.

Run:  python scripts/make_sample_doc.py
Output: data/sample_discharge.docx
"""
from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.shared import Pt

OUT = Path(__file__).resolve().parent.parent / "data" / "sample_discharge.docx"


def main() -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading("Pediatric Discharge Summary", level=1)
    doc.add_paragraph("Patient: {patient_name}    DOB: {dob}    MRN: {mrn}")
    doc.add_paragraph("Discharge Date: {discharge_date}    Discharging Provider: [PROVIDER]")

    doc.add_heading("Diagnosis", level=2)
    doc.add_paragraph(
        "Acute otitis media of the right ear. Patient also has a history of mild persistent asthma."
    )

    doc.add_heading("Medications", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Medication"
    hdr[1].text = "Dose"
    hdr[2].text = "Frequency"
    hdr[3].text = "Notes"

    for med in [
        ("Amoxicillin", "400 mg", "twice daily for 10 days", "Take with food."),
        ("Tylenol", "10 mg/kg", "every 4 hours as needed", "For fever above 38.0 °C."),
        ("Albuterol", "2 puffs", "every 4-6 hours as needed", "Use spacer."),
    ]:
        row = table.add_row().cells
        for i, val in enumerate(med):
            row[i].text = val

    doc.add_heading("Activity", level=2)
    doc.add_paragraph(
        "Your child may return to school in 24 hours after starting Amoxicillin if fever-free. "
        "Avoid swimming or getting water in the right ear for 7 days."
    )

    doc.add_heading("When to Call the Doctor", level=2)
    doc.add_paragraph(
        "Call [CLINIC_PHONE] right away if: fever above 39.0 °C lasts more than 24 hours, "
        "your child has trouble breathing, severe ear pain, drainage of pus or blood from the ear, "
        "or your child is not drinking fluids."
    )

    doc.add_heading("Follow-Up", level=2)
    doc.add_paragraph(
        "Schedule a follow-up visit with {pcp_name} in 2 weeks. Bring this discharge summary and the medication bottles."
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
